from mpmath import mp
from docx import Document
from docx.shared import Cm
from process_img import *
from docx.shared import RGBColor
from docx.shared import Pt

if __name__ == '__main__':
    dps = 14417 - 4*89
    mp.dps=14417 - 4*89 #m* n -1
    text=str(mp.pi)
    margin = 1.27

    doc = Document()

    # Define the text and formatting
    font_size = Pt(6)
    font_colors = process()
    font: "Cambria"
    
    # Add paragraphs with different font sizes and colors
    p= doc.add_paragraph()
    
    for i in range(len(font_colors)):
        digit = text[i]
        color = font_colors[i]
        run = p.add_run(digit)
        run.font.color.rgb = RGBColor(color[0], color[1], color[2])
        run.font.size = font_size
    
    # Add credits to the footer
    credits = "Source code: https://github.com/annieqqa/happyPiDay\nImage source: https://www.vecteezy.com/free-vector/happy-pi-day"
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = credits

    # sections = doc.sections
    # for section in sections:
    section.top_margin = Cm(margin)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)

    # Save the Word document
    doc.save("output/HappyPiDay.docx")
    print("done!")